#!/usr/bin/env python3
import docx
from pathlib import Path

def read_docx(file_path):
    """Read and display the content of a .docx file."""
    doc = docx.Document(file_path)
    
    print("=" * 80)
    print("华中科技大学科研管理平台设计方案")
    print("=" * 80)
    print()
    
    # Read paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            # Check if it's a heading
            if para.style.name.startswith('Heading'):
                print(f"\n{'#' * int(para.style.name[-1])} {para.text}")
            else:
                print(para.text)
            print()
    
    # Read tables if any
    if doc.tables:
        print("\n" + "=" * 80)
        print("表格内容:")
        print("=" * 80)
        for i, table in enumerate(doc.tables):
            print(f"\n表格 {i+1}:")
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                print(" | ".join(row_data))
            print()

if __name__ == "__main__":
    docx_path = Path("/Users/eric/projects5/fanny3/doc/华中科技大学科研管理平台设计方案.docx")
    read_docx(docx_path)